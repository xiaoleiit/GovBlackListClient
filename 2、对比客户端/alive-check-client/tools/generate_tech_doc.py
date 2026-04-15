from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


PROJECT_NAME = "政务局黑名单信息对比客户端"
PROJECT_ALIAS = "模块标识：alive-check-client"
VERSION = "v1.0.0"
DOC_DATE = "2026-04-15"
TARGET_READER = "技术开发人员、联调人员、部署维护人员"
CLIENT_ROOT = Path(__file__).resolve().parents[1]
WORKSPACE_ROOT = Path(__file__).resolve().parents[3]
OUTPUT_PATH = WORKSPACE_ROOT / "5、开发相关文档" / "技术文档" / f"{PROJECT_NAME}技术文档.docx"

FONT_YAHEI = "微软雅黑"
FONT_CONSOLAS = "Consolas"
COLOR_BLUE = "1F4E79"
COLOR_LIGHT_GRAY = "F2F2F2"
COLOR_CODE_BG = "F6F8FA"


def set_doc_update_fields(document: Document) -> None:
    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def set_font_attributes(
    target,
    name: str,
    size: Optional[float] = None,
    bold: Optional[bool] = None,
    color: Optional[str] = None,
) -> None:
    target.font.name = name
    if size is not None:
        target.font.size = Pt(size)
    if bold is not None:
        target.font.bold = bold
    if color:
        target.font.color.rgb = RGBColor.from_string(color)

    element = target._element
    if hasattr(element, "get_or_add_rPr"):
        rpr = element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), name)
        rfonts.set(qn("w:hAnsi"), name)
        rfonts.set(qn("w:eastAsia"), name)


def style_run(run, name: str = FONT_YAHEI, size: float = 11, bold: bool = False, color: Optional[str] = None):
    set_font_attributes(run, name, size=size, bold=bold, color=color)


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: Optional[str] = None,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
    font_size: float = 10.5,
    font_name: str = FONT_YAHEI,
) -> None:
    cell.text = text
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(3)
        for run in paragraph.runs:
            style_run(run, name=font_name, size=font_size, bold=bold, color=color)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        set_cell_text(
            header_cells[index],
            header,
            bold=True,
            color="FFFFFF",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            font_size=10.5,
        )
        shade_cell(header_cells[index], COLOR_BLUE)

    for row_index, row_data in enumerate(rows, start=1):
        row_cells = table.add_row().cells
        row_fill = COLOR_LIGHT_GRAY if row_index % 2 == 0 else None
        for col_index, value in enumerate(row_data):
            set_cell_text(row_cells[col_index], value)
            if row_fill:
                shade_cell(row_cells[col_index], row_fill)

    document.add_paragraph()


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.3)
    paragraph.paragraph_format.right_indent = Cm(0.3)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    set_paragraph_shading(paragraph, COLOR_CODE_BG)
    run = paragraph.add_run(code)
    style_run(run, name=FONT_CONSOLAS, size=9.5)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.35
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    style_run(run, size=11)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        if not paragraph.runs:
            run = paragraph.add_run(item)
            style_run(run, size=11)
            continue
        for run in paragraph.runs:
            style_run(run, size=11)
        run = paragraph.add_run(item if not paragraph.runs else "")
        if run.text:
            style_run(run, size=11)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(3)
        if not paragraph.runs:
            run = paragraph.add_run(item)
            style_run(run, size=11)
            continue
        for run in paragraph.runs:
            style_run(run, size=11)
        run = paragraph.add_run(item if not paragraph.runs else "")
        if run.text:
            style_run(run, size=11)


def add_heading(document: Document, text: str, level: int) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        if level == 1:
            style_run(run, size=18, bold=True, color=COLOR_BLUE)
        elif level == 2:
            style_run(run, size=14, bold=True, color=COLOR_BLUE)
        else:
            style_run(run, size=12, bold=True, color=COLOR_BLUE)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    set_font_attributes(normal, FONT_YAHEI, size=11, bold=False)

    for style_name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        style = document.styles[style_name]
        set_font_attributes(style, FONT_YAHEI, size=size, bold=True, color=COLOR_BLUE)

    title = document.styles["Title"]
    set_font_attributes(title, FONT_YAHEI, size=24, bold=True, color=COLOR_BLUE)

    subtitle = document.styles["Subtitle"]
    set_font_attributes(subtitle, FONT_YAHEI, size=14, bold=False)

    for style_name in ("List Bullet", "List Number"):
        style = document.styles[style_name]
        set_font_attributes(style, FONT_YAHEI, size=11, bold=False)


def add_cover(document: Document) -> None:
    for _ in range(6):
        document.add_paragraph()

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(PROJECT_NAME)
    style_run(title_run, size=24, bold=True, color=COLOR_BLUE)

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("技术文档")
    style_run(subtitle_run, size=16, bold=True, color=COLOR_BLUE)

    document.add_paragraph()

    meta_lines = [
        f"版本：{VERSION}",
        f"日期：{DOC_DATE}",
        f"目标读者：{TARGET_READER}",
        PROJECT_ALIAS,
    ]
    for line in meta_lines:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(10)
        run = paragraph.add_run(line)
        style_run(run, size=12)

    document.add_page_break()


def add_history_page(document: Document) -> None:
    add_heading(document, "文档历史记录", 1)
    add_table(
        document,
        ["日期", "项目", "版本", "更新记录"],
        [[DOC_DATE, PROJECT_NAME, VERSION, "统一项目名称，补充联调、部署、排障、测试与安全章节。"]],
    )
    document.add_page_break()


def add_toc_page(document: Document) -> None:
    add_heading(document, "目录", 1)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)

    run_begin = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin._r.append(fld_begin)

    run_instr = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    run_instr._r.append(instr)

    run_sep = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep._r.append(fld_sep)

    text_run = paragraph.add_run("打开文档后如目录未自动刷新，请在 Word 中右键目录并选择“更新域”。")
    style_run(text_run, size=11)

    run_end = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end._r.append(fld_end)

    document.add_page_break()


def get_workspace_tree() -> str:
    return """政务局黑名单信息对比/
├── 1、模拟服务端/
│   └── unified-alive-auth/
├── 2、对比客户端/
│   └── alive-check-client/
├── 3、生存待校验文件ZJPV/
├── 4、生存已校验文件ZJPC/
└── 5、开发相关文档/
    ├── 技术文档/
    ├── 接口文档/
    └── README.md"""


def get_client_tree() -> str:
    return """alive-check-client/
├── main.py
├── README.md
├── requirements.txt
├── tools/
│   └── generate_tech_doc.py
└── src/
    ├── __init__.py
    ├── core/
    │   ├── __init__.py
    │   ├── api_client.py
    │   ├── config_manager.py
    │   ├── crypto.py
    │   ├── file_parser.py
    │   └── processor.py
    ├── gui/
    │   ├── __init__.py
    │   ├── app.py
    │   ├── config_panel.py
    │   ├── process_panel.py
    │   └── result_dialog.py
    └── utils/
        ├── __init__.py
        └── logger.py"""


def add_overview_section(document: Document) -> None:
    add_heading(document, "1. 项目概述", 1)

    add_heading(document, "1.1 背景", 2)
    add_body_paragraph(
        document,
        "本项目面向政务场景中的黑名单信息核验流程，提供一个桌面端客户端以批量读取 ZJPV 待校验文件，"
        "调用生存认证接口完成人员状态比对，并输出仅包含已故人员的 ZJPC 文件。项目采用 Python + Tkinter 实现，"
        "目标是在不引入浏览器或复杂运行环境的前提下，为业务人员提供稳定、可操作、可追踪的批处理工具。",
    )

    add_heading(document, "1.2 目标", 2)
    add_body_paragraph(
        document,
        "系统通过统一配置管理、国密加密封装、接口调用、文件解析和 GUI 交互，降低政务业务对接复杂度，"
        "保证文件输入输出格式一致，并为后续打包发布、现场部署和问题排查提供标准化交付材料。",
    )

    add_heading(document, "1.3 技术栈", 2)
    add_table(
        document,
        ["技术项", "版本/组件", "用途说明"],
        [
            ["开发语言", "Python 3.9+", "承载桌面客户端、文件处理、网络调用与脚本化交付。"],
            ["GUI 框架", "Tkinter / ttk", "实现配置页、处理页、结果对话框和窗口切换。"],
            ["密码算法", "gmssl（SM3 / SM4）", "实现请求加密、响应解密与签名生成。"],
            ["HTTP 客户端", "requests", "调用 Token 接口和生存比对接口。"],
            ["文档工具", "python-docx", "生成项目技术文档。"],
            ["打包工具", "PyInstaller", "生成桌面分发包，降低部署环境要求。"],
        ],
    )


def add_architecture_section(document: Document) -> None:
    add_heading(document, "2. 系统架构", 1)

    add_heading(document, "2.1 系统边界", 2)
    add_table(
        document,
        ["对象", "位置", "职责说明"],
        [
            ["模拟服务端", "1、模拟服务端/unified-alive-auth/", "提供 Token 和生存比对接口，支撑本地联调与功能验证。"],
            ["客户端", "2、对比客户端/alive-check-client/", "桌面端批处理工具，负责配置录入、文件处理、接口调用与结果输出。"],
            ["样例输入", "3、生存待校验文件ZJPV/", "存放 ZJPV 输入样例，用于回归和演示。"],
            ["样例输出", "4、生存已校验文件ZJPC/", "存放 ZJPC 结果样例，用于核对输出格式。"],
            ["项目文档", "5、开发相关文档/", "归档技术文档、接口文档及联调依据。"],
        ],
    )

    add_heading(document, "2.2 目录概览", 2)
    add_code_block(document, get_workspace_tree())
    add_code_block(document, get_client_tree())

    add_heading(document, "2.3 模块职责", 2)
    add_table(
        document,
        ["模块", "核心文件", "职责说明"],
        [
            ["入口层", "main.py", "初始化运行环境并启动 AliveCheckApp 主窗口。"],
            ["配置层", "src/core/config_manager.py", "维护默认配置、读取/保存本地配置文件、提供配置校验基础能力。"],
            ["安全层", "src/core/crypto.py", "封装 SM4 加解密与 SM3 签名，屏蔽 gmssl 细节。"],
            ["接口层", "src/core/api_client.py", "封装请求体、Token 获取、响应解密和业务接口调用。"],
            ["文件层", "src/core/file_parser.py", "解析 ZJPV 文件、生成 ZJPC 文件、扫描输入目录。"],
            ["流程层", "src/core/processor.py", "负责批量遍历、分批调用接口、聚合已故人员和输出统计结果。"],
            ["界面层", "src/gui/*.py", "实现配置录入、目录选择、进度显示、结果提示和页面导航。"],
            ["日志预留层", "src/utils/logger.py", "提供 logging 封装，便于后续接入文件日志或排障日志。"],
        ],
    )

    add_heading(document, "2.4 处理流程", 2)
    add_numbered(
        document,
        [
            "用户在配置页面填写 appId、appSecret、Token 接口、比对接口和结果保存目录，并可先执行连通性测试。",
            "配置确认后进入处理页，选择包含 ZJPV 文件的业务目录，GUI 在后台线程启动批处理流程。",
            "BatchProcessor 扫描目录、解析所有 ZJPV 文件并汇总人员信息，同时保留源文件头部信息用于输出继承。",
            "ApiClient 先调用 Token 接口获取令牌，再按每批 100 人执行生存比对，所有请求均走 SM4 加密与 SM3 签名。",
            "处理进度通过 after 回调刷新到 Tk 主线程，界面会展示已处理人数、已故人数和当前批次来源文件。",
            "处理结果中 aliveStatus 为 DEAD 的记录会被汇总，并由 ZJPCWriter 写入目标目录下的“2.生存已校验文件ZJPC”子目录。",
            "任一环节出现异常时，GUI 会弹出错误提示并保留日志文本，便于联调和排障。",
        ],
    )


def add_core_module_section(document: Document) -> None:
    add_heading(document, "3. 核心模块详解", 1)

    add_heading(document, "3.1 config_manager 模块", 2)
    add_body_paragraph(
        document,
        "ConfigManager 负责管理客户端本地配置。默认将配置文件写入用户目录 ~/.alive-check/config.json，"
        "并在对象初始化时自动创建目录、加载已保存参数。该模块通过 DEFAULT_CONFIG 提供兜底值，"
        "保证首次运行也能带出默认接口地址和默认保存目录。",
    )
    add_table(
        document,
        ["要点", "实现说明"],
        [
            ["核心类", "ConfigManager"],
            ["默认配置项", "appId、appSecret、tokenUrl、compareUrl、saveDir"],
            ["关键方法", "load()、save()、get()、set()、update()、is_configured()、clear()"],
            ["设计价值", "将 GUI 表单与持久化细节解耦，便于后续替换配置来源或增加字段。"],
        ],
    )

    add_heading(document, "3.2 crypto 模块", 2)
    add_body_paragraph(
        document,
        "crypto 模块是所有接口安全封装的基础能力。模块对 appSecret 做 16 字节截断或补零处理，"
        "使用 gmssl 实现 SM4 ECB 模式加解密和 PKCS7 填充，并提供 SM3(secret + encryptData) 的签名规则实现。"
        "当前加密输出默认使用 Hex，也兼容 Base64 解密输入。",
    )
    add_table(
        document,
        ["要点", "实现说明"],
        [
            ["核心函数", "encrypt_data()、decrypt_data()、sign_data()"],
            ["加密模式", "SM4 ECB + PKCS7 Padding"],
            ["签名规则", "SM3(appSecret + encryptData)"],
            ["兼容性处理", "解密时先尝试 Base64，再回退 Hex；兼容 gmssl 返回 bytes 或 list 的两种形式。"],
        ],
    )

    add_heading(document, "3.3 api_client 模块", 2)
    add_body_paragraph(
        document,
        "ApiClient 面向接口调用流程进行封装。其内部通过 _build_request_body 构造统一报文壳，"
        "通过 _decrypt_response 统一处理 respCode 与 encryptData 的校验，并对 Token 生命周期进行缓存。"
        "业务代码只需调用 generate_token() 和 alive_compare()，无需直接接触加解密细节。",
    )
    add_table(
        document,
        ["要点", "实现说明"],
        [
            ["核心类", "ApiClient"],
            ["关键方法", "_build_request_body()、_decrypt_response()、generate_token()、ensure_token()、alive_compare()、test_connection()"],
            ["请求头", "Content-Type: application/json; charset=utf-8；比对接口额外携带 token 请求头。"],
            ["Token 管理", "基于 createTime + expiresIn * 1000 计算过期时间，并在失效后自动刷新。"],
            ["当前约束", "requests 调用未显式设置 timeout，生产部署建议补充超时与重试策略。"],
        ],
    )

    add_heading(document, "3.4 file_parser 模块", 2)
    add_body_paragraph(
        document,
        "file_parser 模块负责处理批量文件的输入与输出标准。ZJPVParser 读取输入文件、抽取首两行头信息并解析第 3 行开始的人员记录；"
        "ZJPCWriter 在输出时复用源文件头部格式，仅更新记录数，并将已故人员记录写入结果文件。find_zjpv_files() 提供目录扫描能力。",
    )
    add_table(
        document,
        ["要点", "实现说明"],
        [
            ["输入解析", "文件名前缀为 ZJPV，人员记录按“姓名 + 空格 + 身份证号”拆分。"],
            ["输出生成", "文件名前缀为 ZJPC，默认按当前日期生成 10 位序号文件名。"],
            ["格式继承", "输出文件的 line1/line2 优先继承首个源文件头信息，仅重写记录数字段。"],
            ["输出目录", "saveDir/2.生存已校验文件ZJPC/"],
        ],
    )

    add_heading(document, "3.5 processor 模块", 2)
    add_body_paragraph(
        document,
        "BatchProcessor 是系统的业务编排核心。它先收敛目录中所有人员数据，再以每批 100 人的固定批次调用比对接口。"
        "处理过程中会维护 processed_count、deceased_count 和 total_count，并通过 progress_callback 与 log_callback 向 GUI 实时反馈。"
        "当前实现还会携带当前批次来源文件名，避免进度区长期显示为空。",
    )
    add_table(
        document,
        ["要点", "实现说明"],
        [
            ["核心类", "BatchProcessor"],
            ["关键方法", "process_directory()、stop()"],
            ["批处理策略", "先批量生成 Token，再循环调用 alive_compare()，请求间 sleep 0.1 秒降低速率。"],
            ["结果对象", "返回 total_files、total_persons、processed_persons、deceased_persons、output_file。"],
        ],
    )


def add_gui_section(document: Document) -> None:
    add_heading(document, "4. GUI 模块详解", 1)

    add_heading(document, "4.1 主应用 app.py", 2)
    add_body_paragraph(
        document,
        "AliveCheckApp 负责 Tk 根窗口初始化、配置管理器加载和页面切换。应用启动时默认进入配置页，"
        "通过 on_config_next、on_process_back、on_process_complete 三个回调串联页面流转，并在 run() 中执行窗口居中和 mainloop。"
        "主窗口标题统一为“政务局黑名单信息对比客户端”。",
    )

    add_heading(document, "4.2 配置页 config_panel.py", 2)
    add_body_paragraph(
        document,
        "ConfigPanel 承担配置录入与预检查职责，界面包含 appId、appSecret、Token 接口、比对接口和保存目录。"
        "用户可通过文件夹选择器设置保存位置，点击“测试连接”时会即时创建 ApiClient 调用 test_connection()，"
        "并使用状态标签与消息框反馈结果；点击“下一步”时可将配置写入本地配置文件。",
    )
    add_table(
        document,
        ["交互项", "说明"],
        [
            ["输入控件", "5 个文本或目录输入项，覆盖接口接入的最小必要参数。"],
            ["校验逻辑", "validate_config() 逐项检查空值，避免进入处理页后才暴露配置错误。"],
            ["持久化策略", "save_config_var 为真时，调用 ConfigManager.update() + save() 保存配置。"],
        ],
    )

    add_heading(document, "4.3 处理页 process_panel.py", 2)
    add_body_paragraph(
        document,
        "ProcessPanel 面向批处理操作设计，包含目录选择、进度条、统计信息、滚动日志和开始/停止/返回按钮。"
        "真正的接口调用发生在后台线程 _process_thread 中，避免阻塞主线程；所有 UI 更新通过 after(0, ...) 回到 Tk 主线程执行，"
        "这是该模块保证界面可响应性的关键实现点。",
    )
    add_table(
        document,
        ["交互项", "说明"],
        [
            ["开始比对", "校验目录存在后，重置进度与日志，禁用返回和开始按钮。"],
            ["停止处理", "调用 BatchProcessor.stop() 并恢复按钮状态。"],
            ["进度反馈", "update_progress() 更新百分比、已处理人数、已故人数和当前批次来源文件。"],
            ["异常处理", "_on_error() 记录日志并以消息框提示错误。"],
        ],
    )

    add_heading(document, "4.4 结果对话框 result_dialog.py", 2)
    add_body_paragraph(
        document,
        "ResultDialog 在批处理完成后以模态对话框展示结果摘要，包括总处理人数、已故人数和输出文件名。"
        "当存在输出文件时，可根据不同操作系统调用 explorer、open 或 xdg-open 打开目录或定位文件，"
        "兼顾 Windows、macOS 和 Linux 的桌面使用场景。",
    )


def add_integration_section(document: Document) -> None:
    add_heading(document, "5. 配置与接口联调", 1)

    add_heading(document, "5.1 配置项说明", 2)
    add_table(
        document,
        ["配置项", "是否必填", "说明", "错误影响"],
        [
            ["appId", "是", "开放接口应用标识。", "错误会导致签名校验或身份识别失败。"],
            ["appSecret", "是", "开放接口应用密钥，同时参与 SM4 加密和 SM3 签名。", "错误会导致加解密和签名全部失败。"],
            ["tokenUrl", "是", "Token 接口地址。", "不可达时无法获取 token，所有业务调用失败。"],
            ["compareUrl", "是", "生存比对接口地址。", "不可达时无法执行状态比对。"],
            ["saveDir", "是", "结果保存根目录。", "不可写时无法生成 ZJPC 输出文件。"],
        ],
    )

    add_heading(document, "5.2 通用请求体结构", 2)
    add_table(
        document,
        ["字段", "类型", "说明"],
        [
            ["appId", "String", "应用唯一标识，来自本地配置。"],
            ["encryptType", "String", "固定值 SM4。"],
            ["signType", "String", "固定值 SM3。"],
            ["encryptData", "String", "对业务报文 JSON 执行 SM4 加密后的密文。"],
            ["sign", "String", "对 appSecret + encryptData 做 SM3 计算得到的签名。"],
            ["timestamp", "String", "毫秒级时间戳。"],
            ["version", "String", "当前实现固定为 1.0.0。"],
        ],
    )

    add_heading(document, "5.3 Token 接口业务载荷", 2)
    add_table(
        document,
        ["字段", "类型", "说明"],
        [
            ["appId", "String", "应用 ID。"],
            ["appSecret", "String", "应用密钥，同时用作 SM4 密钥原始输入。"],
        ],
    )

    add_heading(document, "5.4 生存比对接口业务载荷", 2)
    add_table(
        document,
        ["字段", "类型", "说明"],
        [
            ["username", "String", "人员姓名。"],
            ["idcard", "String", "身份证号。"],
        ],
    )
    add_body_paragraph(
        document,
        "比对接口以人员数组作为业务载荷，BatchProcessor 当前固定按每批 100 人组织请求。接口返回结果中，"
        "系统仅将 aliveStatus 等于 DEAD 的记录纳入输出文件，其余状态仅用于统计。",
    )

    add_heading(document, "5.5 联调检查清单", 2)
    add_bullets(
        document,
        [
            "确认开放接口后台已分配 appId、appSecret，并完成调用端 IP 白名单配置。",
            "确认 SM4 输出格式是 Hex 还是 Base64；当前实现默认发送 Hex，解密阶段兼容两者。",
            "确认 SM4 密钥处理方式为 appSecret 截断到 16 字节，不足部分补 0；若服务端规则不同需同步调整 crypto.py。",
            "确认签名规则为 appSecret + encryptData 后做 SM3，而不是对原始业务报文签名。",
            "确认 Token 接口返回 createTime、expiresIn、token，且比对接口请求头携带 token。",
            "确认响应报文要求 respCode == 0 且 encryptData 非空，否则客户端会直接抛出异常。",
            "联调时优先使用单条测试数据验证加密、签名和解密，再放量到批处理。"],
    )

    add_heading(document, "5.6 请求与响应示例", 2)
    add_code_block(
        document,
        '{\n'
        '  "appId": "testApp",\n'
        '  "encryptType": "SM4",\n'
        '  "signType": "SM3",\n'
        '  "encryptData": "0123abcd...",\n'
        '  "sign": "abcd1234...",\n'
        '  "timestamp": "1685411525703",\n'
        '  "version": "1.0.0"\n'
        '}\n\n'
        'Token 解密后示例：\n'
        '{\n'
        '  "token": "mock-token-001",\n'
        '  "createTime": 1685411525703,\n'
        '  "expiresIn": 7200\n'
        '}\n\n'
        '比对结果解密后示例：\n'
        '[\n'
        '  {\n'
        '    "username": "张三",\n'
        '    "idcard": "220101199001011234",\n'
        '    "aliveStatus": "DEAD"\n'
        '  }\n'
        ']',
    )


def add_deploy_section(document: Document) -> None:
    add_heading(document, "6. 运行、部署与发版", 1)

    add_heading(document, "6.1 安装依赖", 2)
    add_code_block(document, "python3 -m pip install -r requirements.txt")

    add_heading(document, "6.2 运行命令", 2)
    add_code_block(document, "cd alive-check-client\npython3 main.py")

    add_heading(document, "6.3 打包命令", 2)
    add_code_block(
        document,
        "# Windows\npyinstaller --onefile --windowed main.py\n\n"
        "# macOS\npyinstaller --onefile --windowed --osx-bundle-identifier com.example.alivecheck main.py\n\n"
        "# Linux\npyinstaller --onefile main.py",
    )

    add_heading(document, "6.4 部署说明", 2)
    add_bullets(
        document,
        [
            "运行环境要求 Python 3.9+，并确保可以访问 Token 与生存比对接口。",
            "如采用打包交付，建议在目标操作系统上重新执行 PyInstaller，避免跨平台二进制兼容问题。",
            "配置文件默认落盘到 ~/.alive-check/config.json，现场部署时需关注用户目录权限。",
            "业务输出文件默认写入配置目录下的“2.生存已校验文件ZJPC”子目录，应保证磁盘写权限。",
            "正式部署前需先验证网络连通性、接口白名单和结果目录的可写权限。",
        ],
    )

    add_heading(document, "6.5 文档生成", 2)
    add_code_block(
        document,
        "python3 -m pip install python-docx\npython3 tools/generate_tech_doc.py",
    )


def add_troubleshooting_section(document: Document) -> None:
    add_heading(document, "7. 异常处理与排障", 1)

    add_heading(document, "7.1 常见异常场景", 2)
    add_table(
        document,
        ["场景", "可能原因", "建议处理方式"],
        [
            ["测试连接失败", "appId/appSecret 错误，或 Token 接口不可达。", "先检查接口地址和网络，再核对签名规则与密钥。"],
            ["接口返回 respCode 非 0", "业务校验失败、签名失败、白名单未配置或 token 异常。", "记录 respMsg，回到服务端日志和接口文档核对规则。"],
            ["响应解密失败", "SM4 输出格式、密钥截断规则或填充方式不一致。", "优先用单条样例验证加解密结果，与服务端逐项比对。"],
            ["未找到 ZJPV 文件", "输入目录为空，或文件名前缀不符合要求。", "确认目录下存在以 ZJPV 开头的业务文件。"],
            ["处理完成但无输出文件", "接口返回结果中没有 aliveStatus=DEAD 的记录。", "核对测试数据与服务端返回值，确认是否存在已故人员。"],
            ["输出文件写入失败", "saveDir 无写权限，或磁盘路径不存在。", "检查目录权限、磁盘空间和目标路径是否正确。"],
        ],
    )

    add_heading(document, "7.2 建议排查顺序", 2)
    add_numbered(
        document,
        [
            "先验证客户端配置项是否填写完整，尤其是 appId、appSecret、tokenUrl、compareUrl。",
            "再使用“测试连接”功能验证 Token 接口是否可用。",
            "若 Token 正常，再用单条人员数据验证比对接口、token 请求头和返回解密是否正确。",
            "若单条请求正常，再检查批量处理逻辑、输入文件格式和输出目录权限。",
            "若问题仍未定位，结合服务端日志、接口文档和客户端日志文本进行交叉比对。",
        ],
    )


def add_file_spec_section(document: Document) -> None:
    add_heading(document, "8. 文件格式规范", 1)

    add_heading(document, "8.1 ZJPV 输入格式", 2)
    add_table(
        document,
        ["项", "规则说明"],
        [
            ["文件名", "前缀 ZJPV + 日期 YYYYMMDD + 10 位序号，例如 ZJPV202604140000000001。"],
            ["第 1 行", "文件标识，解析时原样保留，用于输出文件继承。"],
            ["第 2 行", "交易头记录，前 6 位为记录数，后续 20 位保留域。"],
            ["第 3 行起", "按“姓名 + 空格 + 身份证号”组织，每行一条人员记录。"],
        ],
    )

    add_heading(document, "8.2 ZJPC 输出格式", 2)
    add_table(
        document,
        ["项", "规则说明"],
        [
            ["文件名", "前缀 ZJPC + 当前日期 YYYYMMDD + 10 位序号。"],
            ["输出范围", "仅写入 aliveStatus == DEAD 的人员。"],
            ["第 1 行", "优先继承首个源文件的文件标识。"],
            ["第 2 行", "继承源文件保留域，仅更新前 6 位记录数为实际已故人数。"],
            ["记录区", "姓名截断或补齐到 10 位，身份证号截断或补齐到 20 位。"],
        ],
    )

    add_heading(document, "8.3 边界与异常规则", 2)
    add_bullets(
        document,
        [
            "空行会被自动跳过，不会写入人员列表。",
            "文件名前缀不是 ZJPV 的文件不会被纳入批处理。",
            "姓名和身份证号按空白字符拆分，缺少任一字段的行会被忽略。",
            "输出文件会按照字段定长规则进行截断或补齐，需提前评估超长姓名的业务可接受性。",
            "当前实现不会主动校验身份证号合法性，若业务要求更严，需在接口前增加格式校验。",
        ],
    )

    add_heading(document, "8.4 示例", 2)
    add_code_block(
        document,
        "ZJPV 输入示例\n"
        "01\n"
        "000002FFFFFFFFFFFFFFFFFFFF\n"
        "张三 220101199001011234\n"
        "李四 220101199202023456\n\n"
        "ZJPC 输出示例\n"
        "01\n"
        "000001FFFFFFFFFFFFFFFFFFFF\n"
        "张三      220101199001011234  ",
    )


def add_test_section(document: Document) -> None:
    add_heading(document, "9. 测试与验收", 1)

    add_heading(document, "9.1 建议测试用例", 2)
    add_table(
        document,
        ["测试场景", "输入条件", "预期结果"],
        [
            ["单文件正常流程", "1 个合法 ZJPV 文件，接口可用。", "成功处理并按结果生成 ZJPC 文件。"],
            ["多文件批处理", "多个 ZJPV 文件，总人数跨多个文件。", "统计总人数正确，进度随批次更新。"],
            ["空目录", "选择无 ZJPV 文件目录。", "日志提示未找到 ZJPV 文件，不生成输出。"],
            ["无已故人员", "接口返回均非 DEAD。", "处理完成但不生成输出文件。"],
            ["多批次请求", "总人数超过 100。", "自动分批调用接口，最终结果汇总正确。"],
            ["Token 过期重取", "模拟 token 失效。", "ensure_token() 触发重新获取 token。"],
            ["异常接口响应", "respCode 非 0 或 encryptData 为空。", "界面展示错误信息，处理终止。"],
            ["边界数据", "超长姓名、空行、缺字段记录。", "非法或异常记录不影响程序整体运行。"],
        ],
    )

    add_heading(document, "9.2 验收标准", 2)
    add_bullets(
        document,
        [
            "客户端能够成功连接模拟服务端并完成一次完整比对流程。",
            "输出文件格式符合 ZJPC 规范，头信息继承和记录数更新正确。",
            "界面可正确展示处理进度、处理人数、已故人数和当前批次来源文件。",
            "技术文档能够独立支持安装、运行、联调、排障和交付整理。",
        ],
    )


def add_security_section(document: Document) -> None:
    add_heading(document, "10. 安全与运维注意事项", 1)
    add_bullets(
        document,
        [
            "appSecret 当前以明文形式保存在 ~/.alive-check/config.json 中，正式环境建议限制配置目录权限。",
            "联调和排障过程中不要在日志、截图或工单中暴露 appSecret、token 和完整身份证号。",
            "对接真实接口前需确认调用端 IP 白名单已生效，否则可能出现协议层成功但业务校验失败。",
            "建议将样例文件与正式业务文件分开存放，避免测试数据混入生产目录。",
            "若后续增加文件日志，需优先设计脱敏策略，至少屏蔽 appSecret、token 和敏感身份信息。",
        ],
    )


def add_appendix(document: Document) -> None:
    add_heading(document, "附录", 1)

    add_heading(document, "附录A 配置示例", 2)
    add_code_block(
        document,
        '{\n'
        '    "appId": "APP001",\n'
        '    "appSecret": "secret001abcdef",\n'
        '    "tokenUrl": "http://localhost:8080/openapi/stoken",\n'
        '    "compareUrl": "http://localhost:8080/openapi/aliveCompare",\n'
        '    "saveDir": "/Users/xxx/Documents/生存比对结果"\n'
        '}',
    )

    add_heading(document, "附录B 常见问题", 2)
    add_table(
        document,
        ["问题", "可能原因", "建议处理方式"],
        [
            ["测试连接失败", "appId/appSecret 填写错误，或 Token 接口不可达。", "先检查接口地址和网络，再核对签名规则与密钥。"],
            ["进度区域当前文件为空", "旧版处理器未记录批次来源文件。", "升级到当前版本后重新运行。"],
            ["处理完成但无输出文件", "接口返回结果中没有 aliveStatus=DEAD 的记录。", "核对测试数据与服务端返回值，确认是否存在已故人员。"],
            ["文档目录未显示页码", "Word 未刷新 TOC 域。", "打开文档后右键目录并执行“更新域”。"],
        ],
    )


def build_document() -> Document:
    document = Document()
    set_doc_update_fields(document)
    set_page_layout(document)
    configure_styles(document)

    document.core_properties.title = PROJECT_NAME
    document.core_properties.subject = "技术文档"
    document.core_properties.author = "OpenAI Codex"
    document.core_properties.comments = f"基于当前 {CLIENT_ROOT.name} 项目代码自动整理生成"

    add_cover(document)
    add_history_page(document)
    add_toc_page(document)
    add_overview_section(document)
    add_architecture_section(document)
    add_core_module_section(document)
    add_gui_section(document)
    add_integration_section(document)
    add_deploy_section(document)
    add_troubleshooting_section(document)
    add_file_spec_section(document)
    add_test_section(document)
    add_security_section(document)
    add_appendix(document)

    return document


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(f"技术文档已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
